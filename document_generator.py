from docx import Document

doc = Document("ELP-Form.docx")

doc.tables[0].add_row()
doc.tables[0].rows[5].cells[0].add_paragraph("shit") # 여기부터 2행

doc.save("test.docx")